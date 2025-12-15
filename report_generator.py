"""
Gerador de Relatórios em PDF e Word com marca GAV
Gran Garden Resort - Sistema de Gestão de Obras
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Tentar importar Image, mas não é crítico
try:
    from reportlab.platypus import Image
    HAS_IMAGE_SUPPORT = True
except ImportError:
    HAS_IMAGE_SUPPORT = False


class ReportGenerator:
    """Classe para gerar relatórios profissionais em PDF e Word"""
    
    def __init__(self, logo_path='public/assets/gran-garden-resort.jpg'):
        """
        Inicializa o gerador de relatórios
        
        Args:
            logo_path: Caminho para o logo da GAV
        """
        self.logo_path = logo_path
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Configura estilos customizados para os relatórios"""
        # Estilo para título principal
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Estilo para subtítulo
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#334155'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        ))
        
        # Estilo para corpo de texto
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#475569'),
            spaceAfter=6,
            alignment=TA_JUSTIFY
        ))
        
        # Estilo para rodapé
        self.styles.add(ParagraphStyle(
            name='CustomFooter',
            parent=self.styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#64748b'),
            alignment=TA_CENTER
        ))
    
    def generate_pdf(self, data, output_path):
        """
        Gera relatório em PDF
        
        Args:
            data: Dicionário com os dados do relatório
            output_path: Caminho do arquivo de saída
        """
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Container para elementos do documento
        story = []
        
        # Adicionar logo (se disponível)
        if HAS_IMAGE_SUPPORT and os.path.exists(self.logo_path):
            try:
                logo = Image(self.logo_path, width=8*cm, height=4*cm)
                logo.hAlign = 'CENTER'
                story.append(logo)
                story.append(Spacer(1, 1*cm))
            except Exception as e:
                print(f"Aviso: Não foi possível adicionar logo: {e}")
        
        # Título principal
        title = Paragraph(
            f"GRAN GARDEN RESORT<br/>RELATÓRIO {data.get('type', '').upper()}",
            self.styles['CustomTitle']
        )
        story.append(title)
        story.append(Spacer(1, 0.5*cm))
        
        # Linha separadora
        story.append(self._create_separator())
        story.append(Spacer(1, 0.5*cm))
        
        # Informações do projeto
        info_data = [
            ['Projeto:', data.get('project', '-')],
            ['Período:', data.get('period_label', '-')],
            ['Data de Geração:', datetime.now().strftime('%d/%m/%Y %H:%M')],
        ]
        
        info_table = Table(info_data, colWidths=[5*cm, 12*cm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#334155')),
            ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#475569')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 1*cm))
        
        # Linha separadora
        story.append(self._create_separator())
        story.append(Spacer(1, 0.5*cm))
        
        # Resumo Executivo
        story.append(Paragraph('RESUMO EXECUTIVO', self.styles['CustomSubtitle']))
        story.append(Spacer(1, 0.3*cm))
        
        summary_data = [
            ['Métrica', 'Valor'],
            ['Total de Tarefas', str(data.get('tasks', 0))],
            ['Tarefas Concluídas', str(data.get('completedTasks', 0))],
            ['Percentual de Conclusão', f"{data.get('completion_percentage', 0):.2f}%"],
        ]
        
        summary_table = Table(summary_data, colWidths=[12*cm, 5*cm])
        summary_table.setStyle(TableStyle([
            # Cabeçalho
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            # Corpo
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#334155')),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 11),
            ('ALIGN', (0, 1), (0, -1), 'LEFT'),
            ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
            # Bordas e grid
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 1*cm))
        
        # Observações adicionais
        if data.get('includeCharts'):
            story.append(Paragraph(
                '<b>Gráficos:</b> Incluídos neste relatório',
                self.styles['CustomBody']
            ))
        
        if data.get('includePhotos'):
            story.append(Paragraph(
                '<b>Registro Fotográfico:</b> Incluído neste relatório',
                self.styles['CustomBody']
            ))
        
        # Rodapé
        story.append(Spacer(1, 2*cm))
        story.append(self._create_separator())
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(
            'Relatório gerado automaticamente pelo Gran Garden Resort<br/>'
            'Sistema de Gestão de Obras v2.0 | GAV Resorts<br/>'
            f'© {datetime.now().year} - Todos os direitos reservados',
            self.styles['CustomFooter']
        ))
        
        # Construir PDF
        doc.build(story)
        print(f"✅ PDF gerado: {output_path}")
    
    def generate_word(self, data, output_path):
        """
        Gera relatório em Word
        
        Args:
            data: Dicionário com os dados do relatório
            output_path: Caminho do arquivo de saída
        """
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar logo (se disponível)
        if os.path.exists(self.logo_path):
            try:
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(self.logo_path, width=Inches(4))
            except Exception as e:
                print(f"Aviso: Não foi possível adicionar logo: {e}")
        
        # Título principal
        title = doc.add_heading('GRAN GARDEN RESORT', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(30, 64, 175)
        
        subtitle = doc.add_heading(f'RELATÓRIO {data.get("type", "").upper()}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = RGBColor(51, 65, 85)
        
        doc.add_paragraph()  # Espaço
        
        # Informações do projeto
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_cells = [
            ('Projeto:', data.get('project', '-')),
            ('Período:', data.get('period_label', '-')),
            ('Data de Geração:', datetime.now().strftime('%d/%m/%Y %H:%M')),
        ]
        
        for i, (label, value) in enumerate(info_cells):
            row = info_table.rows[i]
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            label_cell.text = label
            label_run = label_cell.paragraphs[0].runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            
            value_cell.text = value
            value_run = value_cell.paragraphs[0].runs[0]
            value_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Espaço
        
        # Resumo Executivo
        doc.add_heading('RESUMO EXECUTIVO', level=2)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Medium Grid 3 Accent 1'
        
        # Cabeçalho
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'Métrica'
        header_cells[1].text = 'Valor'
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dados
        summary_data = [
            ('Total de Tarefas', str(data.get('tasks', 0))),
            ('Tarefas Concluídas', str(data.get('completedTasks', 0))),
            ('Percentual de Conclusão', f"{data.get('completion_percentage', 0):.2f}%"),
        ]
        
        for i, (metric, value) in enumerate(summary_data, start=1):
            row = summary_table.rows[i]
            row.cells[0].text = metric
            row.cells[1].text = value
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
            
            # Alinhar valor à direita
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Espaço
        
        # Observações adicionais
        if data.get('includeCharts'):
            p = doc.add_paragraph()
            p.add_run('Gráficos: ').bold = True
            p.add_run('Incluídos neste relatório')
        
        if data.get('includePhotos'):
            p = doc.add_paragraph()
            p.add_run('Registro Fotográfico: ').bold = True
            p.add_run('Incluído neste relatório')
        
        # Rodapé
        doc.add_page_break()
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f'Relatório gerado automaticamente pelo Gran Garden Resort\n'
            f'Sistema de Gestão de Obras v2.0 | GAV Resorts\n'
            f'© {datetime.now().year} - Todos os direitos reservados'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)
        
        # Salvar documento
        doc.save(output_path)
        print(f"✅ Word gerado: {output_path}")
    
    def _create_separator(self):
        """Cria uma linha separadora"""
        return Table([['']], colWidths=[17*cm], rowHeights=[2])


def generate_report_from_json(json_data, format='pdf', output_dir='relatorios'):
    """
    Função auxiliar para gerar relatório a partir de dados JSON
    
    Args:
        json_data: Dicionário com os dados do relatório
        format: 'pdf' ou 'word'
        output_dir: Diretório de saída
    """
    # Criar diretório se não existir
    os.makedirs(output_dir, exist_ok=True)
    
    # Calcular percentual de conclusão
    tasks = json_data.get('tasks', 0)
    completed = json_data.get('completedTasks', 0)
    completion_percentage = (completed / tasks * 100) if tasks > 0 else 0
    
    json_data['completion_percentage'] = completion_percentage
    
    # Nome do arquivo
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_type = json_data.get('type', 'geral')
    
    generator = ReportGenerator()
    
    if format.lower() == 'pdf':
        output_path = os.path.join(output_dir, f'relatorio_{report_type}_{timestamp}.pdf')
        generator.generate_pdf(json_data, output_path)
    else:
        output_path = os.path.join(output_dir, f'relatorio_{report_type}_{timestamp}.docx')
        generator.generate_word(json_data, output_path)
    
    return output_path


# Exemplo de uso
if __name__ == '__main__':
    # Dados de teste
    test_data = {
        'type': 'geral',
        'project': 'Gran Garden Resort - Blocos A, B e C',
        'period': 'last-month',
        'period_label': 'Último Mês',
        'tasks': 150,
        'completedTasks': 120,
        'includeCharts': True,
        'includePhotos': True,
        'generatedAt': datetime.now().isoformat()
    }
    
    print("="*80)
    print("TESTE DO GERADOR DE RELATÓRIOS")
    print("="*80)
    
    # Gerar PDF
    pdf_path = generate_report_from_json(test_data, format='pdf')
    print(f"\nPDF gerado: {pdf_path}")
    
    # Gerar Word
    word_path = generate_report_from_json(test_data, format='word')
    print(f"Word gerado: {word_path}")
    
    print("\n" + "="*80)
    print("TESTE CONCLUÍDO COM SUCESSO!")
    print("="*80)
